"""
cv_parser.py
Extracts raw text from uploaded CV files (PDF or DOCX).
"""

import io
import pdfplumber
import docx


class UnsupportedFileTypeError(Exception):
    """Raised when a file isn't a .pdf or .docx"""
    pass


class EmptyDocumentError(Exception):
    """Raised when a file parses successfully but yields no usable text."""
    pass


def _extract_pdf_text(file_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, since many CVs use table layouts
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs).strip()


def extract_text(uploaded_file) -> str:
    """
    Extract text from a Streamlit UploadedFile object.
    Raises UnsupportedFileTypeError or EmptyDocumentError on failure.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        text = _extract_pdf_text(file_bytes)
    elif filename.endswith(".docx"):
        text = _extract_docx_text(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"'{uploaded_file.name}' is not a supported file type. "
            "Please upload a .pdf or .docx file."
        )

    if not text or len(text.strip()) < 20:
        raise EmptyDocumentError(
            f"'{uploaded_file.name}' appears to be empty, image-only, or "
            "unreadable. Try a text-based PDF/DOCX instead of a scanned image."
        )

    return text
